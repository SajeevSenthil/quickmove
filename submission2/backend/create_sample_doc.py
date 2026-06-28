"""
Generate sample customer requirement .docx files for testing.
  cd submission2/backend
  python create_sample_doc.py
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_doc(output_path: str):
    doc = Document()

    title = doc.add_heading("Customer Relocation Requirement", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_heading("Personal Details", level=1)
    details = [
        ("Customer Name", "Arjun Mehta"),
        ("Contact Email", "arjun.mehta@techcorp.in"),
        ("Contact Phone", "+91 98765 43210"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_heading("Relocation Details", level=1)
    relocation = [
        ("Destination City", "Bengaluru"),
        ("Current City", "Pune"),
        ("Target Move Date", "1st August 2026"),
        ("Reason for Relocation", "Job transfer to ITPL campus"),
    ]
    for label, value in relocation:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_heading("Apartment Requirements", level=1)
    apartment = [
        ("Apartment Type", "2BHK"),
        ("Monthly Rent Budget", "₹25,000 – ₹35,000"),
        ("Preferred Localities", "Whitefield, Marathahalli, ITPL area, Brookfield"),
        ("Furnished Status", "Semi-furnished (basic appliances provided)"),
        ("Parking", "Yes, 1 car parking required"),
        ("Floor Preference", "2nd floor or above"),
    ]
    for label, value in apartment:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_heading("Additional Requirements", level=1)
    special = [
        ("Pets", "No pets"),
        ("Security", "Gated community / 24x7 security preferred"),
        ("Power Backup", "Power backup required for work from home"),
        ("Internet", "Fibre internet connectivity (Airtel/Jio preferred)"),
        ("Office Distance", "Within 5 km of ITPL, Whitefield"),
    ]
    for label, value in special:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_heading("Timeline", level=1)
    doc.add_paragraph(
        "The customer needs to move by 1st August 2026. Property shortlisting should be "
        "completed by 20th July 2026 to allow sufficient time for lease signing and logistics."
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Sample document saved to: {output_path}")


if __name__ == "__main__":
    create_sample_doc(os.path.join(os.path.dirname(__file__), "..", "sample_docs", "arjun_mehta_requirement.docx"))
