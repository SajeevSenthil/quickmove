import io
import json
import re
from docx import Document
from google import genai
from ..config import settings

EXTRACTION_PROMPT = """You are an AI assistant for QuickMove, a relocation company in India.

Extract structured relocation requirements from the following customer document.

Document:
{text}

Return a JSON object with these exact fields:
{{
  "name": "customer full name as string",
  "city": "destination city as string",
  "budget_min": minimum monthly rent in INR as integer (e.g. 20000),
  "budget_max": maximum monthly rent in INR as integer (e.g. 35000),
  "preferred_regions": ["list", "of", "preferred", "localities/areas"],
  "apartment_type": "1BHK or 2BHK or 3BHK or Studio etc as string",
  "furnished_status": "furnished or semi-furnished or unfurnished",
  "parking_required": true or false,
  "pets_allowed": true or false,
  "office_location": "office area or address for commute reference as string",
  "move_date": "YYYY-MM-DD format or null",
  "special_requirements": ["any", "special", "needs", "as", "list"]
}}

Rules:
- Return ONLY valid JSON, no markdown fences, no explanation text
- If a field is not mentioned, use null for strings or empty array for lists
- Budgets must be integers in INR per month (convert "25k" to 25000)
- If only one budget number is given, set both min and max to that value
- For move_date: if only a day/month is given with no year, use the current year 2026
- If no customer name is found, set name to "Unknown Customer"
"""

DEMO_EXTRACTION = {
    "name": "Arjun Mehta",
    "city": "Bengaluru",
    "budget_min": 25000,
    "budget_max": 35000,
    "preferred_regions": ["Whitefield", "Marathahalli", "Brookfield"],
    "apartment_type": "2BHK",
    "furnished_status": "semi-furnished",
    "parking_required": True,
    "pets_allowed": False,
    "office_location": "ITPL, Whitefield",
    "move_date": "2026-08-01",
    "special_requirements": ["24/7 security", "backup power"],
}


def extract_text_from_bytes(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


async def extract_customer_requirements(text: str) -> dict:
    if settings.DEMO_MODE or not settings.GEMINI_API_KEY:
        return DEMO_EXTRACTION

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    prompt = EXTRACTION_PROMPT.format(text=text)

    response = await client.aio.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )
    raw = response.text.strip()

    # Strip markdown code fences if the model wraps its output
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise ValueError(f"Gemini returned non-JSON response: {raw[:300]}")
